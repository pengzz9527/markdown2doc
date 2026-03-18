import streamlit as st
import duckdb
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from bs4 import BeautifulSoup
import re
import os
import io
import base64
import hashlib
from datetime import datetime
from pathlib import Path
import sys

# 检查并导入 markdown 库
try:
    import markdown
    # 测试 markdown 是否可用
    test_html = markdown.markdown("# Test")
    if not test_html or test_html.strip() == "":
        raise ImportError("markdown 库返回空内容")
    MARKDOWN_AVAILABLE = True
    st.sidebar.success(f"✅ Markdown 库已加载")
except Exception as e:
    st.error(f"❌ Markdown 库导入失败: {e}")
    st.info("请运行: pip uninstall markdown && pip install markdown")
    MARKDOWN_AVAILABLE = False
    st.stop()

# 配置页面
st.set_page_config(
    page_title="Markdown 转换工具",
    page_icon="📝",
    layout="wide",
    initial_sidebar_state="expanded"
)

# 自定义CSS样式
st.markdown("""
<style>
    .main-header {
        font-size: 2.5rem;
        font-weight: bold;
        color: #1f77b4;
        margin-bottom: 1rem;
    }
    .preview-box {
        border: 1px solid #e0e0e0;
        border-radius: 8px;
        padding: 20px;
        background-color: #ffffff;
        max-height: 600px;
        overflow-y: auto;
        min-height: 200px;
        color: #333333 !important;
    }
    .preview-box * {
        color: #333333 !important;
    }
    .preview-box h1, .preview-box h2, .preview-box h3 {
        color: #2c3e50 !important;
        margin-top: 24px;
        margin-bottom: 16px;
    }
    .preview-box p {
        color: #333333 !important;
        margin: 10px 0;
        line-height: 1.6;
    }
    .preview-box ul, .preview-box ol {
        color: #333333 !important;
        margin: 10px 0;
        padding-left: 20px;
    }
    .preview-box li {
        color: #333333 !important;
        margin: 5px 0;
    }
    .stButton>button {
        width: 100%;
    }
    .download-link {
        display: inline-block;
        padding: 8px 16px;
        background-color: #1f77b4;
        color: white;
        text-decoration: none;
        border-radius: 4px;
        margin: 5px 0;
        font-size: 14px;
    }
</style>
""", unsafe_allow_html=True)

# 初始化 DuckDB 数据库
@st.cache_resource
def init_db():
    """初始化 DuckDB 数据库，创建转换历史表"""
    db_path = "markdown_history.db"
    conn = duckdb.connect(db_path)
    conn.execute("""
              CREATE SEQUENCE IF NOT EXISTS doc2_seq START 1;
              """)
    conn.execute("""
        CREATE TABLE IF NOT EXISTS conversions (
            id INTEGER  DEFAULT nextval('doc2_seq'),
            timestamp TIMESTAMP DEFAULT CURRENT_TIMESTAMP,
            title VARCHAR,
            content_hash VARCHAR UNIQUE,
            markdown_content TEXT,
            word_blob BLOB,
            html_blob BLOB,
            file_size_kb INTEGER
        )
    """)
    
    return conn

def save_to_history(conn, title, markdown_content, word_bytes, html_bytes):
    """保存转换记录到 DuckDB"""
    if not markdown_content:
        return False, "内容为空"
    
    content_hash = hashlib.md5(markdown_content.encode()).hexdigest()
    file_size = len(markdown_content) // 1024
    
    try:
        conn.execute("""
            INSERT INTO conversions (title, content_hash, markdown_content, word_blob, html_blob, file_size_kb)
            VALUES (?, ?, ?, ?, ?, ?)
        """, [title, content_hash, markdown_content, word_bytes, html_bytes, file_size])
        return True, "新记录"
    except duckdb.ConstraintException:
        conn.execute("""
            UPDATE conversions 
            SET timestamp = CURRENT_TIMESTAMP, word_blob = ?, html_blob = ?
            WHERE content_hash = ?
        """, [word_bytes, html_bytes, content_hash])
        return True, "已更新"

def get_history(conn, limit=20):
    """获取转换历史"""
    try:
        return conn.execute("""
            SELECT id, timestamp, title, content_hash, file_size_kb 
            FROM conversions 
            ORDER BY timestamp DESC 
            LIMIT ?
        """, [limit]).fetchdf()
    except Exception as e:
        st.error(f"读取历史记录失败: {e}")
        return None

def get_conversion_by_id(conn, conv_id):
    """根据 ID 获取具体转换记录"""
    try:
        result = conn.execute("""
            SELECT title, word_blob, html_blob, markdown_content 
            FROM conversions 
            WHERE id = ?
        """, [conv_id]).fetchone()
        return result
    except Exception as e:
        st.error(f"获取记录失败: {e}")
        return None

def delete_history_item(conn, conv_id):
    """删除历史记录"""
    try:
        conn.execute("DELETE FROM conversions WHERE id = ?", [conv_id])
        return True
    except Exception as e:
        st.error(f"删除失败: {e}")
        return False

class MarkdownToWordConverter:
    def __init__(self):
        self.doc = Document()
        self.doc.styles['Normal'].font.name = 'Microsoft YaHei'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
    
    def convert(self, markdown_text):
        """将 Markdown 转换为 Word 文档"""
        if not markdown_text:
            raise ValueError("Markdown 内容不能为空")
        
        # 使用基础扩展，避免扩展加载失败
        try:
            html = markdown.markdown(
                markdown_text, 
                extensions=['tables', 'fenced_code', 'toc', 'nl2br']
            )
        except Exception as e:
            # 如果扩展加载失败，使用基础模式
            st.warning(f"扩展加载失败，使用基础模式: {e}")
            html = markdown.markdown(markdown_text)
        
        if not html:
            raise ValueError("Markdown 转换返回空内容")
        
        soup = BeautifulSoup(html, 'html.parser')
        
        for element in soup.find_all(['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'pre', 'table', 'ul', 'ol', 'blockquote']):
            self._process_element(element)
        
        return self.doc
    
    def _process_element(self, element):
        """处理单个 HTML 元素"""
        if not element or not element.name:
            return
            
        tag = element.name
        
        try:
            if tag.startswith('h'):
                level = int(tag[1])
                self._add_heading(element.get_text(), level)
            elif tag == 'p':
                self._add_paragraph(element)
            elif tag == 'pre':
                self._add_code_block(element)
            elif tag == 'table':
                self._add_table(element)
            elif tag in ['ul', 'ol']:
                self._add_list(element, ordered=(tag == 'ol'))
            elif tag == 'blockquote':
                self._add_blockquote(element)
        except Exception as e:
            st.warning(f"处理元素 {tag} 时出错: {e}")
    
    def _add_heading(self, text, level):
        """添加标题"""
        if not text:
            return
        heading = self.doc.add_heading(text, level=level)
        for run in heading.runs:
            run.font.name = 'Microsoft YaHei'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            if level == 1:
                run.font.color.rgb = RGBColor(31, 119, 180)
    
    def _add_paragraph(self, element):
        """添加段落（支持行内格式）"""
        p = self.doc.add_paragraph()
        
        for content in element.contents:
            try:
                if content.name == 'strong':
                    run = p.add_run(content.get_text())
                    run.bold = True
                elif content.name == 'em':
                    run = p.add_run(content.get_text())
                    run.italic = True
                elif content.name == 'code':
                    run = p.add_run(content.get_text())
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                elif content.name == 'a':
                    run = p.add_run(content.get_text())
                    run.font.color.rgb = RGBColor(0, 0, 255)
                    run.underline = True
                elif isinstance(content, str):
                    p.add_run(content)
                
                for run in p.runs:
                    run.font.name = 'Microsoft YaHei'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Microsoft YaHei')
            except Exception:
                try:
                    p.add_run(str(content))
                except:
                    pass
    
    def _add_code_block(self, element):
        """添加代码块"""
        code_text = element.get_text() or ""
        p = self.doc.add_paragraph()
        run = p.add_run(code_text)
        run.font.name = 'Courier New'
        run.font.size = Pt(9)
        
        try:
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            shading_elm = parse_xml(r'<w:shd {} w:fill="F5F5F5"/>'.format(nsdecls('w')))
            p._p.get_or_add_pPr().append(shading_elm)
        except:
            pass
            
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)
    
    def _add_table(self, element):
        """添加表格"""
        rows = element.find_all('tr')
        if not rows:
            return
        
        try:
            header_cells = rows[0].find_all(['th', 'td'])
            num_cols = len(header_cells)
            num_rows = len(rows)
            
            table = self.doc.add_table(rows=num_rows, cols=num_cols)
            table.style = 'Light Grid Accent 1'
            
            for i, row in enumerate(rows):
                cells = row.find_all(['th', 'td'])
                for j, cell in enumerate(cells):
                    if j < num_cols:
                        table.rows[i].cells[j].text = cell.get_text() or ""
                        if i == 0 or cell.name == 'th':
                            for paragraph in table.rows[i].cells[j].paragraphs:
                                for run in paragraph.runs:
                                    run.bold = True
        except Exception as e:
            st.warning(f"表格处理出错: {e}")
    
    def _add_list(self, element, ordered=False):
        """添加列表"""
        try:
            items = element.find_all('li', recursive=False)
            for i, item in enumerate(items):
                p = self.doc.add_paragraph(style='List Number' if ordered else 'List Bullet')
                text = item.get_text() or ""
                p.add_run(text)
                
                nested = item.find(['ul', 'ol'], recursive=False)
                if nested:
                    self._add_list(nested, ordered=(nested.name == 'ol'))
        except Exception as e:
            st.warning(f"列表处理出错: {e}")
    
    def _add_blockquote(self, element):
        """添加引用块"""
        try:
            text = element.get_text() or ""
            p = self.doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(text)
            run.italic = True
            run.font.color.rgb = RGBColor(100, 100, 100)
            
            from docx.oxml import parse_xml
            from docx.oxml.ns import nsdecls
            pPr = p._p.get_or_add_pPr()
            pBdr = parse_xml(r'<w:pBdr {}><w:left w:val="single" w:sz="24" w:space="4" w:color="CCCCCC"/></w:pBdr>'.format(nsdecls('w')))
            pPr.append(pBdr)
        except Exception as e:
            st.warning(f"引用块处理出错: {e}")
    
    def save_to_bytes(self):
        """保存文档到字节流"""
        doc_io = io.BytesIO()
        self.doc.save(doc_io)
        doc_io.seek(0)
        return doc_io.getvalue()

def create_html_download(markdown_text):
    """创建完整的 HTML 文件（带样式）"""
    if not markdown_text:
        return b""
    
    try:
        # 尝试使用扩展，失败则使用基础模式
        try:
            html_content = markdown.markdown(
                markdown_text,
                extensions=['tables', 'fenced_code', 'toc', 'nl2br']
            )
        except:
            html_content = markdown.markdown(markdown_text)
        
        if not html_content:
            return b""
        
        styled_html = f"""<!DOCTYPE html>
<html lang="zh-CN">
<head>
    <meta charset="UTF-8">
    <meta name="viewport" content="width=device-width, initial-scale=1.0">
    <title>Markdown 导出</title>
    <style>
        body {{
            font-family: -apple-system, BlinkMacSystemFont, "Segoe UI", Roboto, "Helvetica Neue", Arial, sans-serif;
            line-height: 1.6;
            color: #333;
            max-width: 900px;
            margin: 0 auto;
            padding: 40px 20px;
            background-color: #fff;
        }}
        h1, h2, h3, h4, h5, h6 {{
            margin-top: 24px;
            margin-bottom: 16px;
            font-weight: 600;
            line-height: 1.25;
            color: #2c3e50;
        }}
        h1 {{ font-size: 2em; border-bottom: 2px solid #eaecef; padding-bottom: 0.3em; }}
        h2 {{ font-size: 1.5em; border-bottom: 1px solid #eaecef; padding-bottom: 0.3em; }}
        code {{
            background-color: #f6f8fa;
            padding: 0.2em 0.4em;
            border-radius: 3px;
            font-family: SFMono-Regular, Consolas, "Liberation Mono", Menlo, Courier, monospace;
            font-size: 85%;
        }}
        pre {{
            background-color: #f6f8fa;
            padding: 16px;
            overflow: auto;
            border-radius: 6px;
            line-height: 1.45;
        }}
        pre code {{
            background-color: transparent;
            padding: 0;
        }}
        table {{
            border-collapse: collapse;
            width: 100%;
            margin: 16px 0;
        }}
        th, td {{
            border: 1px solid #dfe2e5;
            padding: 6px 13px;
        }}
        th {{
            background-color: #f6f8fa;
            font-weight: 600;
        }}
        tr:nth-child(2n) {{
            background-color: #f6f8fa;
        }}
        blockquote {{
            margin: 0;
            padding: 0 1em;
            color: #6a737d;
            border-left: 0.25em solid #dfe2e5;
        }}
        img {{
            max-width: 100%;
            height: auto;
        }}
    </style>
</head>
<body>
    {html_content}
</body>
</html>"""
        return styled_html.encode('utf-8')
    except Exception as e:
        st.error(f"HTML 生成失败: {e}")
        return b""

def get_download_link(file_bytes, filename, mime_type):
    """生成下载链接"""
    if not file_bytes:
        return ""
    try:
        b64 = base64.b64encode(file_bytes).decode()
        href = f'<a href="data:{mime_type};base64,{b64}" download="{filename}" class="download-link">⬇️ 下载 {filename}</a>'
        return href
    except Exception as e:
        return f"生成下载链接失败: {e}"

def render_markdown_safe(text):
    """安全地渲染 Markdown 预览"""
    if not text or not text.strip():
        return "<p style='color: #999;'>暂无内容...</p>"
    
    try:
        # 首先尝试带扩展的渲染
        try:
            html_content = markdown.markdown(
                text,
                extensions=['tables', 'fenced_code', 'nl2br']
            )
        except Exception as ext_error:
            # 如果扩展失败，使用基础渲染
            st.sidebar.warning(f"扩展加载失败: {ext_error}，使用基础模式")
            html_content = markdown.markdown(text)
        
        # 检查渲染结果
        if not html_content or html_content.strip() == "":
            return f"<pre style='color: red;'>渲染结果为空，原始内容：\n{text[:200]}...</pre>"
        
        # 清理 HTML，确保没有空标签导致显示问题
        html_content = html_content.replace('<p></p>', '')
        html_content = html_content.replace('<p> </p>', '')
        
        return html_content
        
    except Exception as e:
        error_msg = str(e)
        st.error(f"预览生成失败: {error_msg}")
        import traceback
        st.code(traceback.format_exc())
        return f"<pre style='color: red;'>{text[:500]}...</pre>"

# 主应用逻辑
def main():
    st.markdown('<div class="main-header">📝 Markdown 转换工具</div>', unsafe_allow_html=True)
    st.markdown("将 Markdown 转换为精美的 Word 文档或 HTML 网页")
    
    # 初始化数据库
    try:
        conn = init_db()
    except Exception as e:
        st.error(f"数据库初始化失败: {e}")
        return
    
    # 初始化 session state
    if 'markdown_content' not in st.session_state:
        st.session_state.markdown_content = ""
    if 'loaded_from_history' not in st.session_state:
        st.session_state.loaded_from_history = False
    
    # 创建两列布局
    col1, col2 = st.columns([3, 2])
    
    with col1:
        st.subheader("📥 输入 Markdown")
        
        # 输入方式选择
        input_method = st.radio(
            "选择输入方式：",
            ["直接输入", "上传文件"],
            horizontal=True,
            key="input_method"
        )
        
        # 处理输入
        current_content = st.session_state.markdown_content
        
        if input_method == "直接输入":
            # 如果有从历史加载的内容，使用它
            default_value = current_content if st.session_state.loaded_from_history else ""
            
            new_content = st.text_area(
                "在此粘贴 Markdown 内容：",
                value=default_value,
                height=400,
                placeholder="# 标题\n\n开始编写你的 Markdown 内容...\n\n## 特性支持\n- **粗体** 和 *斜体*\n- `行内代码`\n- 代码块\n- [链接](https://example.com)\n- 表格\n- 列表",
                key="markdown_input"
            )
            # 更新 session state
            if new_content != current_content:
                st.session_state.markdown_content = new_content
                st.session_state.loaded_from_history = False
                
        else:  # 上传文件
            uploaded_file = st.file_uploader("上传 Markdown 文件", type=['md', 'markdown', 'txt'], key="file_uploader")
            if uploaded_file is not None:
                try:
                    file_content = uploaded_file.read().decode('utf-8')
                    st.session_state.markdown_content = file_content
                    st.success(f"✅ 已加载文件：{uploaded_file.name} ({len(file_content)} 字符)")
                except Exception as e:
                    st.error(f"文件读取失败: {e}")
        
        # 获取当前内容（确保不为 None）
        markdown_content = st.session_state.markdown_content or ""
        
        # 操作按钮区域
        if markdown_content.strip():
            st.markdown("---")
            cols = st.columns(3)
            
            with cols[0]:
                convert_word = st.button("📝 转换为 Word", use_container_width=True, type="primary", key="btn_word")
            with cols[1]:
                convert_html = st.button("🌐 转换为 HTML", use_container_width=True, type="primary", key="btn_html")
            with cols[2]:
                save_history_btn = st.button("💾 保存到历史", use_container_width=True, key="btn_save")
            
            # 执行转换
            if convert_word or convert_html or save_history_btn:
                with st.spinner("正在转换..."):
                    try:
                        # 生成标题（取第一行标题或前20字符）
                        title = "未命名文档"
                        lines = markdown_content.strip().split('\n')
                        for line in lines:
                            if line.startswith('#'):
                                title = line.lstrip('#').strip()
                                break
                        if title == "未命名文档":
                            title = markdown_content[:30] + "..." if len(markdown_content) > 30 else markdown_content
                        
                        # 生成 Word
                        converter = MarkdownToWordConverter()
                        converter.convert(markdown_content)
                        word_bytes = converter.save_to_bytes()
                        
                        # 生成 HTML
                        html_bytes = create_html_download(markdown_content)
                        
                        # 保存到历史
                        if save_history_btn or True:  # 总是保存，但只在点击保存按钮时显示提示
                            success, msg = save_to_history(conn, title, markdown_content, word_bytes, html_bytes)
                            if save_history_btn:
                                if success:
                                    st.success(f"✅ 已保存到历史记录！({msg})")
                                    st.balloons()
                                else:
                                    st.error(f"保存失败: {msg}")
                        
                        # 提供下载
                        if convert_word:
                            word_link = get_download_link(word_bytes, f"{title[:20]}.docx", 
                                            "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                            if word_link:
                                st.markdown(word_link, unsafe_allow_html=True)
                            else:
                                st.error("Word 文件生成失败")
                        
                        if convert_html:
                            html_link = get_download_link(html_bytes, f"{title[:20]}.html", "text/html")
                            if html_link:
                                st.markdown(html_link, unsafe_allow_html=True)
                            else:
                                st.error("HTML 文件生成失败")
                                
                    except Exception as e:
                        st.error(f"转换过程出错: {e}")
                        st.exception(e)
        else:
            st.info("👈 请在左侧输入 Markdown 内容或上传文件")
    
    with col2:
        st.subheader("👁️ 实时预览")
        
        # 安全获取内容
        preview_content = st.session_state.markdown_content or ""
        
        if preview_content.strip():
            # 使用安全的渲染函数
            html_preview = render_markdown_safe(preview_content)
            
            # 调试信息（取消注释以查看）
            # with st.expander("Debug - HTML 源码"):
            #     st.code(html_preview[:2000])
            
            # 渲染预览
            preview_container = st.container()
            with preview_container:
                st.markdown('<div class="preview-box">' + html_preview + '</div>', unsafe_allow_html=True)
        else:
            st.info("在左侧输入内容以查看预览")
        
        # 统计信息
        if preview_content:
            st.markdown("---")
            st.subheader("📊 统计")
            col_stat1, col_stat2, col_stat3 = st.columns(3)
            with col_stat1:
                st.metric("字符数", len(preview_content))
            with col_stat2:
                word_count = len(preview_content.split())
                st.metric("词数", word_count)
            with col_stat3:
                line_count = len(preview_content.split('\n'))
                st.metric("行数", line_count)
    
    # 历史记录区域
    st.markdown("---")
    st.subheader("🕐 转换历史")
    
    try:
        history_df = get_history(conn)
        
        if history_df is None or len(history_df) == 0:
            st.info("暂无转换历史")
        else:
            # 搜索功能
            search_term = st.text_input("🔍 搜索历史记录", placeholder="输入关键词搜索...", key="search_history")
            
            if search_term:
                filtered = history_df[history_df['title'].str.contains(search_term, case=False, na=False)]
            else:
                filtered = history_df
            
            if len(filtered) == 0:
                st.info("没有找到匹配的记录")
            else:
                # 显示历史记录
                for _, row in filtered.iterrows():
                    with st.container():
                        cols = st.columns([3, 2, 1, 1])
                        
                        with cols[0]:
                            display_title = row['title'][:40] if len(row['title']) > 40 else row['title']
                            st.markdown(f"**{display_title}**")
                            st.caption(f"{row['timestamp'].strftime('%Y-%m-%d %H:%M')} · {row['file_size_kb']} KB")
                        
                        with cols[1]:
                            # 获取文件数据
                            result = get_conversion_by_id(conn, row['id'])
                            if result:
                                title, word_blob, html_blob, md_content = result
                                
                                # Word 下载
                                word_link = get_download_link(word_blob, f"doc_{row['id']}.docx", 
                                                "application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                                if word_link:
                                    st.markdown(word_link, unsafe_allow_html=True)
                                
                                # HTML 下载
                                html_link = get_download_link(html_blob, f"html_{row['id']}.html", "text/html")
                                if html_link:
                                    st.markdown(html_link, unsafe_allow_html=True)
                            else:
                                st.caption("文件不可用")
                        
                        with cols[2]:
                            # 重新加载按钮
                            if st.button("🔄 加载", key=f"load_{row['id']}"):
                                res = get_conversion_by_id(conn, row['id'])
                                if res and len(res) >= 4:
                                    st.session_state.markdown_content = res[3] or ""
                                    st.session_state.loaded_from_history = True
                                    st.rerun()
                                else:
                                    st.error("无法加载内容")
                        
                        with cols[3]:
                            # 删除按钮
                            if st.button("🗑️ 删除", key=f"del_{row['id']}"):
                                if delete_history_item(conn, row['id']):
                                    st.success("已删除")
                                    st.rerun()
                    
                    st.markdown("---")
    except Exception as e:
        st.error(f"加载历史记录时出错: {e}")
    
    # 页脚
    st.markdown("""
    <div style="text-align: center; color: #666; padding: 20px; margin-top: 40px; border-top: 1px solid #eee;">
        <small>支持标准 Markdown 语法 · 表格 · 代码块 · 自动保存历史记录</small>
    </div>
    """, unsafe_allow_html=True)

if __name__ == "__main__":
    main()